from dotenv import load_dotenv
import os
from openai import OpenAI
from docx import Document

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def generate_text(prompt: str) -> str:
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an elite AI consultant, outbound strategist, automation architect, "
                    "and small-business operator. "
                    "You use strong business judgment. "
                    "You never give generic advice. "
                    "You optimize for a solo entrepreneur selling AI services to small and mid-size businesses. "
                    "You avoid enterprise-company assumptions unless clearly labeled as a translated benchmark. "
                    "You write in practical 2026 business language. "
                    "You focus on revenue, speed, missed opportunities, labor savings, conversion lift, "
                    "retention lift, and operational leverage. "
                    "You are persuasive but realistic."
                )
            },
            {"role": "user", "content": prompt}
        ],
        temperature=0.9
    )
    return response.choices[0].message.content

def add_section(doc: Document, title: str, body: str):
    doc.add_heading(title, level=1)
    for block in body.split("\n\n"):
        cleaned = block.strip()
        if cleaned:
            doc.add_paragraph(cleaned)

def main():
    print("\nGenerating final AI sales document... please wait...\n")

    company = "Hudson Financial Group"
    website = "https://www.hfgventures.net/"

    company_context_prompt = f"""
Research-minded analysis for this business:

Company: {company}
Website: {website}

You are creating a SELLABLE AI opportunity document for a solo entrepreneur.

Hard rules:
- Assume this is an insurance / financial services business unless evidence strongly suggests otherwise.
- Small to mid-size business framing only.
- No enterprise recommendations unless translated down to SMB reality.
- No fluff.
- No basic filler ideas.
- No implementation tutorials.
- Only outcomes, business value, pain points, blind spots, and sellable AI solutions.

Return these sections with strong detail:

1. COMPANY REALITY
- What this company likely sells
- Who it likely serves
- How it likely gets customers
- Where revenue is likely won or lost
- Where time is likely wasted
- What operational friction probably exists

2. INDUSTRY-SPECIFIC BLIND SPOTS AND PAIN POINTS
Give 12 realistic pain points and blind spots.
For each include:
- Problem
- Why it likely exists
- What it likely costs
- How AI fixes it
- Expected outcome

3. HARD IMPACT PROJECTIONS
Use LOW / MID / HIGH scenarios for:
- revenue lift
- lead response speed
- conversion lift
- time saved monthly
- monthly cost savings
- yearly cost savings
- client retention lift

4. 5 URGENT AI FACTS
Facts a business owner in this industry would care about.
Make them strong, specific, and persuasive.

5. WHAT TOP COMPANIES IN THIS INDUSTRY ARE DOING WITH AI
But explain it in a way that is relatable to a small or mid-size operator.

Style:
- direct
- commercially sharp
- specific
- no soft language
"""

    acquisition_prompt = f"""
Company: {company}
Website: {website}

Create the section:
CUSTOMER ACQUISITION AI OPPORTUNITIES

Rules:
- 12 high-quality solutions only
- no filler
- must include at least:
  - 3 AI agents
  - 3 voice agents
  - 3 complex workflow automations
- simple to advanced progression
- only things a solo entrepreneur can realistically sell to a small/mid-size insurance or financial services business

For EACH solution include:
- Solution name
- Type: simple / intermediate / advanced / complex
- What it does
- Business benefit
- Expected impact
- Why it matters specifically for this type of business

Focus on:
- lead capture
- lead qualification
- speed to lead
- appointment booking
- pipeline generation
- missed inbound opportunities
"""

    sales_prompt = f"""
Company: {company}
Website: {website}

Create the section:
SALES AI OPPORTUNITIES

Rules:
- 12 high-quality solutions only
- no filler
- must include at least:
  - 3 AI agents
  - 3 voice agents
  - 3 complex workflow automations
- only realistic, sellable solutions for a small/mid-size insurance or financial services firm

For EACH solution include:
- Solution name
- Type: simple / intermediate / advanced / complex
- What it does
- Business benefit
- Expected impact
- Why it matters specifically for this business type

Focus on:
- lead qualification
- follow-up
- quote follow-up
- appointment setting
- sales call prep
- missed calls
- no-shows
- reactivation
- pipeline acceleration
"""

    marketing_prompt = f"""
Company: {company}
Website: {website}

Create the section:
MARKETING AI OPPORTUNITIES

Rules:
- 12 high-quality solutions only
- no filler
- must include at least:
  - 3 AI agents
  - 3 voice agents
  - 3 complex workflow automations
- must be relevant to a small/mid-size insurance or financial services business

For EACH solution include:
- Solution name
- Type
- What it does
- Business benefit
- Expected impact
- Why it matters for this business

Focus on:
- content creation
- nurture campaigns
- review generation
- local visibility
- paid ad follow-up
- lead magnet workflows
- segmentation
- remarketing
"""

    service_ops_prompt = f"""
Company: {company}
Website: {website}

Create two sections:
1. CUSTOMER SERVICE AND RETENTION AI OPPORTUNITIES
2. OPERATIONS AND ADMIN AI OPPORTUNITIES

For EACH section:
- 12 high-quality solutions only
- no filler
- must include at least:
  - 3 AI agents
  - 3 voice agents
  - 3 complex workflow automations

For EACH solution include:
- Solution name
- Type
- What it does
- Business benefit
- Expected impact
- Why it matters for this business

Focus on:
- policyholder/client follow-up
- renewal reminders
- retention saves
- support triage
- document handling
- appointment reminders
- onboarding
- admin reduction
- internal reporting
- compliance-friendly process support
"""

    workflow_prompt = f"""
Company: {company}
Website: {website}

Create the section:
ADVANCED AI WORKFLOW SYSTEMS

Give 8 truly useful multi-step systems for a small/mid-size insurance or financial services business.

For EACH workflow include:
- Workflow name
- Trigger
- Step-by-step flow
- Where AI is used
- What gets automated
- Business result
- Why this is high leverage

Make them real, such as:
- inbound lead to booked appointment
- missed call to AI response to follow-up
- quote request to qualification to nurture
- dormant lead reactivation
- review request and referral generation
- renewal retention workflow
"""

    sales_toolkit_prompt = f"""
Company: {company}
Website: {website}

Create these sections:

1. WHAT TO SELL FIRST
- 5 strongest initial offers
- prioritize obvious ROI and easiest close
- explain why each is a smart first sale

2. PITCH STRATEGY
- best opening angle
- what pain to lead with
- how to position value
- what not to say
- likely objections
- how to answer them

3. PERSONAL TEXT MESSAGE
- sounds human
- sounds like the sender knows the business owner personally
- not robotic
- not cheesy

4. COLD EMAIL SEQUENCE
- Email 1
- Follow-up 1
- Follow-up 2
- should sound like a real operator, not corporate AI junk

5. TARGET COMPANIES
Find 30 realistic companies to target.
Rules:
- same or similar industry
- similar size to this company or smaller
- small to mid-size only
- no giant brands
- no enterprise firms
- start hyper-local if possible, then broader Texas
- group by:
  - 50-100 employees
  - 101-250 employees
  - 251-500 employees
  - 500-1000 employees
For each company include:
- Company name
- What they do
- Why they fit
- Likely decision-maker title

6. REUSABLE INDUSTRY ANGLE
- a generalized angle that can be reused for similar companies in the same industry
"""

    print("Writing company analysis...")
    company_context = generate_text(company_context_prompt)

    print("Writing acquisition section...")
    acquisition = generate_text(acquisition_prompt)

    print("Writing sales section...")
    sales = generate_text(sales_prompt)

    print("Writing marketing section...")
    marketing = generate_text(marketing_prompt)

    print("Writing service and ops section...")
    service_ops = generate_text(service_ops_prompt)

    print("Writing workflow systems...")
    workflows = generate_text(workflow_prompt)

    print("Writing sales toolkit...")
    sales_toolkit = generate_text(sales_toolkit_prompt)

    doc = Document()
    doc.add_heading(f"{company} - AI Sales System", 0)
    doc.add_paragraph(f"Website: {website}")

    add_section(doc, "Company Analysis", company_context)
    add_section(doc, "Customer Acquisition AI Opportunities", acquisition)
    add_section(doc, "Sales AI Opportunities", sales)
    add_section(doc, "Marketing AI Opportunities", marketing)
    add_section(doc, "Customer Service and Operations AI Opportunities", service_ops)
    add_section(doc, "Advanced AI Workflow Systems", workflows)
    add_section(doc, "Sales Toolkit", sales_toolkit)

    output_file = "hudson_ai_sales_system.docx"
    doc.save(output_file)

    print("\nCreated:")
    print(output_file)

if __name__ == "__main__":
    main()